import io
import json
import os
import time
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode
from urllib.request import Request, urlopen
from uuid import uuid4

from django.conf import settings
from django.db import transaction
from django.db.models import F, Window
from django.db.models.functions import RowNumber
from django.urls import reverse
from docx import Document as DocxDocument

from common.models import Code
from common.onlyoffice import decode_jwt, encode_jwt
from common.project_selection import get_request_user
from common.storage import build_s3_uri, delete_object, delete_object_at_uri, read_bytes_from_uri, save_bytes
from files.models import ProjectFile
from projects.models import ProjectNet, ProjectUserRole

from .models import Document, DocumentApproval, DocumentDetail


DOCUMENT_CODE_SEQUENCE = ("DOC_SRS", "DOC_ITF", "DOC_ARCH", "DOC_ERD", "DOC_DB", "DOC_TS")

APPROVAL_STATUS_SEQUENCE = ("APRV_REQ", "APRV_COM", "APRV_RJT")
PROJECT_ROLE_CODES = ("ROLE_MANAGER", "ROLE_MEMBER")
GENERATION_SESSION_KEY = "docs_initial_generation"
ALLOWED_GENERATION_FILE_CODES = ("FILE_RFP", "FILE_MEETING")
INTERFACE_REFERENCE_DOCUMENT_CODE = "DOC_ITF"
ARCHITECTURE_DOCUMENT_CODE = "DOC_ARCH"
INTERFACE_REFERENCE_ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg"}
INTERFACE_REFERENCE_MAX_FILE_SIZE = 3 * 1024 * 1024
PROGRESS_PENDING = "PRGRS_PENDING"
PROGRESS_PROCESSING = "PRGRS_PROCESSING"
PROGRESS_COMPLETED = "PRGRS_COMPLETED"
PROGRESS_FAILED = "PRGRS_FAILED"
RUNNING_PROGRESS_CODES = (PROGRESS_PENDING, PROGRESS_PROCESSING)
TERMINAL_PROGRESS_CODES = (PROGRESS_COMPLETED, PROGRESS_FAILED)
FASTAPI_GENERATE_TIMEOUT_SECONDS = 10


def _build_empty_generation_state(project):
    return {
        "project_sn": project.sn if project else None,
        "selected_file_ids": [],
        "draft_documents": {},
        "confirmed_documents": {},
        "itf_reference_files": [],
    }


def _normalize_reference_filename(filename):
    return Path(filename or "").name or "reference"


def build_itf_reference_storage_key(project, actor, filename):
    project_key = getattr(project, "sn", "none")
    actor_key = getattr(actor, "sn", "anonymous")
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else "bin"
    return f"itf-references/{project_key}/{actor_key}/{uuid4().hex}.{extension}"


def _legacy_cleanup_path(path_value):
    if not path_value:
        return
    path = Path(path_value)
    for _ in range(3):
        try:
            path.unlink(missing_ok=True)
            return
        except FileNotFoundError:
            return
        except PermissionError:
            try:
                os.chmod(path, 0o666)
            except OSError:
                pass
            time.sleep(0.05)
        except OSError:
            return


def cleanup_generation_itf_reference(reference):
    storage_key = (reference or {}).get("storage_key", "")
    if storage_key:
        delete_object(storage_key)
        return
    _legacy_cleanup_path((reference or {}).get("path", ""))


def cleanup_generation_itf_references(state):
    for reference in state.get("itf_reference_files", []):
        cleanup_generation_itf_reference(reference)


def get_generation_itf_references(state):
    return list(state.get("itf_reference_files", []))


def get_fastapi_base_url():
    return str(getattr(settings, "FASTAPI_BASE_URL", "") or "").rstrip("/")

def get_doc_job_poll_interval_seconds():
    raw_value = getattr(settings, "DOC_JOB_POLL_INTERVAL_SECONDS", 10)
    try:
        interval = int(raw_value)
    except (TypeError, ValueError):
        return 10
    return interval if interval > 0 else 10


def _build_reference_uri(reference):
    path_value = str((reference or {}).get("path", "") or "").strip()
    if path_value:
        return path_value

    storage_key = str((reference or {}).get("storage_key", "") or "").strip()
    if not storage_key:
        return ""

    try:
        return build_s3_uri(storage_key)
    except Exception:
        return storage_key


def add_generation_itf_references(project, actor, state, uploaded_files):
    references = state.setdefault("itf_reference_files", [])
    added_count = 0
    errors = []

    for uploaded_file in uploaded_files:
        if uploaded_file is None or not getattr(uploaded_file, "name", ""):
            continue

        filename = _normalize_reference_filename(uploaded_file.name)
        extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if extension not in INTERFACE_REFERENCE_ALLOWED_EXTENSIONS:
            errors.append("png, jpg, jpeg 이미지 파일만 업로드할 수 있습니다.")
            continue
        if uploaded_file.size > INTERFACE_REFERENCE_MAX_FILE_SIZE:
            errors.append("각 이미지는 3MB 이하만 업로드할 수 있습니다.")
            continue

        token = uuid4().hex
        storage_key = build_itf_reference_storage_key(project, actor, filename)
        content_bytes = b"".join(uploaded_file.chunks())
        content_type = f"image/{'jpeg' if extension == 'jpg' else extension}"
        save_bytes(storage_key, content_bytes, content_type=content_type)

        references.append(
            {
                "token": token,
                "name": filename,
                "size": uploaded_file.size,
                "extension": extension,
                "storage_key": storage_key,
                "path": _build_reference_uri({"storage_key": storage_key}),
            }
        )
        added_count += 1

    state["itf_reference_files"] = references
    return added_count, errors


def remove_generation_itf_reference(state, token):
    if not token:
        return False

    remaining_references = []
    removed = False
    for reference in state.get("itf_reference_files", []):
        if reference.get("token") == token:
            cleanup_generation_itf_reference(reference)
            removed = True
            continue
        remaining_references.append(reference)

    state["itf_reference_files"] = remaining_references
    return removed


def _get_ordered_codes(code_values):
    code_map = Code.objects.in_bulk(code_values, field_name="code")
    return [code_map[code] for code in code_values if code in code_map]


def get_document_type_rows():
    return _get_ordered_codes(DOCUMENT_CODE_SEQUENCE)


def get_document_code_sequence():
    rows = get_document_type_rows()
    if rows:
        return [row.code for row in rows]
    return list(DOCUMENT_CODE_SEQUENCE)


def get_document_type_map():
    return {row.code: row for row in get_document_type_rows()}


def resolve_document_code(raw_code):
    sequence = get_document_code_sequence()
    if not sequence:
        return raw_code or DOCUMENT_CODE_SEQUENCE[0]
    return raw_code if raw_code in sequence else sequence[0]


def get_document_type_choices(*, include_all=False):
    rows = get_document_type_rows()
    choices = [("all", "전체")] if include_all else []
    for row in rows:
        choices.append((row.code, row.name))
    return tuple(choices)


def get_approval_status_choices(*, include_all=False):
    rows = _get_ordered_codes(APPROVAL_STATUS_SEQUENCE)
    choices = [("all", "전체")] if include_all else []
    for row in rows:
        choices.append((row.code, row.name))
    return tuple(choices)


def get_document_label(document_code):
    row = Code.objects.filter(code=document_code).only("name").first()
    return row.name if row else document_code


def get_document_index(document_code):
    sequence = get_document_code_sequence()
    try:
        return sequence.index(document_code)
    except ValueError:
        return 0


def get_previous_document_code(document_code):
    sequence = get_document_code_sequence()
    try:
        index = sequence.index(document_code)
    except ValueError:
        return None
    return None if index == 0 else sequence[index - 1]


def get_actor(request):
    return get_request_user(request)


def get_project_role(project, user):
    if project is None or user is None:
        return None
    role = (
        ProjectUserRole.objects.filter(
            project=project,
            user=user,
            role_id__in=PROJECT_ROLE_CODES,
        )
        .order_by("-role_id")
        .first()
    )
    return role.role_id if role else None


def is_project_manager(project, user):
    return get_project_role(project, user) == "ROLE_MANAGER"


def is_project_participant(project, user):
    return get_project_role(project, user) in PROJECT_ROLE_CODES


def get_document_title(document):
    return f"{document.document_type_id}_v{document.version}.docx"


def build_document_key(document, latest_detail=None):
    if latest_detail is None:
        latest_detail = get_latest_detail(document)
    timestamp = int(latest_detail.created_at.timestamp()) if latest_detail else document.sn
    return f"docs-{document.sn}-v{document.version}-{timestamp}"


def build_docx_bytes(title, body_lines):
    content = DocxDocument()
    content.add_heading(title, level=0)
    for line in body_lines:
        content.add_paragraph(line)
    buffer = io.BytesIO()
    content.save(buffer)
    return buffer.getvalue()


def extract_text_from_docx(binary_content):
    if not binary_content:
        return ""

    try:
        document = DocxDocument(io.BytesIO(binary_content))
    except Exception:
        return ""

    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text).strip()


def build_document_detail_storage_key(project, document_sn, detail_sn):
    return f"document-details/{project.sn}/{document_sn}/{detail_sn}.docx"


def build_document_detail_path(project, document_sn, detail_sn):
    return build_s3_uri(build_document_detail_storage_key(project, document_sn, detail_sn))


def get_document_detail_bytes(detail):
    if detail is None:
        return b""
    detail_path = str(getattr(detail, "path", "") or "").strip()
    if not detail_path:
        raise ValueError("Document detail is missing docs_path.")
    return read_bytes_from_uri(detail_path)


def download_remote_content(url):
    if not url:
        return None
    with urlopen(url, timeout=10) as response:
        return response.read()


def request_force_save(document, *, latest_detail=None, userdata=None):
    document_server_url = settings.ONLYOFFICE_DOCUMENT_SERVER_URL.rstrip("/")
    if not document_server_url:
        raise ValueError("OnlyOffice Document Server URL is not configured.")

    document_key = build_document_key(document, latest_detail=latest_detail)
    command_url = f"{document_server_url}/command?shardkey={document_key}"
    payload = {
        "c": "forcesave",
        "key": document_key,
    }
    if userdata:
        payload["userdata"] = userdata

    if settings.ONLYOFFICE_JWT_SECRET:
        request_payload = {
            "token": encode_jwt(payload, settings.ONLYOFFICE_JWT_SECRET),
        }
    else:
        request_payload = payload

    request = Request(
        command_url,
        data=json.dumps(request_payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urlopen(request, timeout=10) as response:
        return json.loads(response.read().decode("utf-8") or "{}")


def get_latest_detail(document):
    return document.details.filter(is_deleted="N").order_by("-created_at", "-sn").first()


def get_highest_detail_sn(document):
    if document is None:
        return None
    return document.details.filter(is_deleted="N").order_by("-sn").first()


def get_detail_by_sn(document, detail_sn):
    return document.details.filter(sn=detail_sn, is_deleted="N").order_by("-created_at", "-sn").first()


def get_latest_pending_approval(document):
    return (
        DocumentApproval.objects.filter(
            detail__document=document,
            approval_status_id="APRV_REQ",
        )
        .select_related("detail", "created_by", "approval_status")
        .order_by("-created_at", "-approval_sn")
        .first()
    )


def latest_confirmed_document(project, document_code):
    return (
        Document.objects.filter(project=project, document_type_id=document_code)
        .exclude(version="0")
        .select_related("document_type", "created_by", "possession_user")
        .order_by("-created_at", "-sn")
        .first()
    )


def get_document_history_queryset(project, document_code):
    if project is None:
        return Document.objects.none()

    return (
        Document.objects.filter(project=project, document_type_id=document_code)
        .exclude(version="0")
        .annotate(
            version_rank=Window(
                expression=RowNumber(),
                partition_by=[F("document_type_id"), F("version")],
                order_by=[F("created_at").desc(), F("sn").desc()],
            )
        )
        .filter(version_rank=1)
        .select_related("document_type", "created_by", "possession_user")
        .order_by("-created_at", "-sn")
    )


def has_any_confirmed_initial_document(project):
    if project is None:
        return False
    return Document.objects.filter(
        project=project,
        document_type_id__in=get_document_code_sequence(),
    ).exclude(version="0").exists()


def has_all_generated_document_types(project):
    if project is None:
        return False
    generated_count = (
        Document.objects.filter(
            project=project,
            document_type_id__in=get_document_code_sequence(),
        )
        .exclude(version="0")
        .values_list("document_type_id", flat=True)
        .distinct()
        .count()
    )
    return generated_count >= len(get_document_code_sequence())


def can_start_initial_generation(project, user):
    if project is None or user is None or not is_project_manager(project, user):
        return False
    return not has_all_generated_document_types(project)


def has_active_generation_session(state):
    return (
        bool(
            state.get("selected_file_ids")
            or state.get("draft_documents")
            or state.get("confirmed_documents")
            or state.get("itf_reference_files")
        )
        and not is_generation_complete(state)
    )


def can_access_initial_generation(project, user, state):
    if project is None or user is None or not is_project_manager(project, user):
        return False
    return has_active_generation_session(state) or not has_all_generated_document_types(project)


def is_latest_document_for_type(document):
    if document is None:
        return False
    latest_document = (
        Document.objects.filter(
            project=document.project,
            document_type_id=document.document_type_id,
        )
        .order_by("-sn")
        .first()
    )
    return latest_document is not None and latest_document.sn == document.sn


def is_latest_detail_for_document(document, detail):
    if document is None or detail is None:
        return False
    latest_detail = get_highest_detail_sn(document)
    return latest_detail is not None and latest_detail.sn == detail.sn


def _build_fastapi_generate_url():
    base_url = get_fastapi_base_url()
    if not base_url:
        raise ValueError("FastAPI base URL is not configured.")
    return f"{base_url}/generate"


def request_fastapi_generate(payload):
    request = Request(
        _build_fastapi_generate_url(),
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urlopen(request, timeout=FASTAPI_GENERATE_TIMEOUT_SECONDS) as response:
        body = response.read().decode("utf-8") or "{}"
    try:
        return json.loads(body)
    except json.JSONDecodeError:
        return {"raw": body}


def extract_fastapi_error_message(exc):
    if isinstance(exc, HTTPError):
        try:
            body = exc.read().decode("utf-8")
        except Exception:
            body = ""
        if body:
            try:
                payload = json.loads(body)
            except json.JSONDecodeError:
                payload = None
            if isinstance(payload, dict):
                detail = payload.get("detail")
                if isinstance(detail, str) and detail.strip():
                    return f"FastAPI error {exc.code}: {detail.strip()}"
                if isinstance(detail, list) and detail:
                    first_detail = detail[0]
                    if isinstance(first_detail, dict):
                        message = first_detail.get("msg") or first_detail.get("message")
                        if message:
                            return f"FastAPI error {exc.code}: {message}"
        return f"FastAPI error {exc.code}: {exc.reason}"
    if isinstance(exc, URLError):
        reason = getattr(exc, "reason", None)
        return f"FastAPI connection failed: {reason or exc}"
    if isinstance(exc, ValueError):
        return str(exc)
    return str(exc)


def get_generation_reference_uris(state):
    uris = []
    for reference in get_generation_itf_references(state):
        uri = _build_reference_uri(reference)
        if uri:
            uris.append(uri)
    return uris

#### fast api 요청 payload 만들기
def build_generation_request_payload(project, state, document_code, *, update_mode="N", selected_files=None):
    files = selected_files if selected_files is not None else get_generation_selected_files(project, state)
    image_list = get_generation_reference_uris(state) if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE else []
    return {
        "project_sn": project.sn,
        "docs_cd": "SRS",
        # "docs_cd": document_code,
        "udt_yn": update_mode,
        "file_list": [project_file.sn for project_file in files],
        "image_list": image_list,
        "etc": {"debug": True},
        # "etc": {"debug": bool(getattr(settings, "DEBUG", False))},
    }


def build_auto_apply_request_payload(project, document_code, selected_files):
    return {
        "project_sn": project.sn,
        "docs_cd": document_code,
        "udt_yn": "Y",
        "file_list": [project_file.sn for project_file in selected_files],
        "image_list": [],
        "etc": {"debug": bool(getattr(settings, "DEBUG", False))},
    }


def _document_job_queryset(project, document_code, *, initial_only=False):
    queryset = Document.objects.filter(project=project, document_type_id=document_code).select_related(
        "project",
        "document_type",
        "created_by",
        "updated_by",
        "possession_user",
    )
    if initial_only:
        queryset = queryset.filter(version="0")
    return queryset


def get_generation_draft_document(project, state, document_code=None):
    if project is None:
        return None
    target_code = document_code or get_current_generation_code(state)
    if not target_code:
        return None
    target_sn = state.get("draft_documents", {}).get(target_code)
    if not target_sn:
        return None
    return (
        _document_job_queryset(project, target_code, initial_only=True)
        .filter(sn=target_sn)
        .first()
    )


def set_generation_draft_document(state, document):
    if document is None:
        return state
    state.setdefault("draft_documents", {})[document.document_type_id] = document.sn
    return state


def clear_generation_draft_document(state, document_code):
    state.setdefault("draft_documents", {}).pop(document_code, None)
    return state


def find_document_job(
    project,
    document_code,
    *,
    tracking_document_sn=None,
    initial_only=False,
    progress_codes=None,
):
    if project is None:
        return None
    queryset = _document_job_queryset(project, document_code, initial_only=initial_only)
    if progress_codes:
        queryset = queryset.filter(progress_status_id__in=progress_codes)
    if tracking_document_sn:
        tracked_document = queryset.filter(sn=tracking_document_sn).first()
        if tracked_document is not None:
            return tracked_document
    return queryset.order_by("-updated_at", "-created_at", "-sn").first()


def wait_for_document_job(
    project,
    document_code,
    *,
    tracking_document_sn=None,
    initial_only=False,
    timeout_seconds=3,
    interval_seconds=0.25,
):
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        document = find_document_job(
            project,
            document_code,
            tracking_document_sn=tracking_document_sn,
            initial_only=initial_only,
        )
        if document is not None:
            return document
        time.sleep(interval_seconds)
    return find_document_job(
        project,
        document_code,
        tracking_document_sn=tracking_document_sn,
        initial_only=initial_only,
    )


def get_running_initial_document(project, document_code, *, tracking_document_sn=None):
    return find_document_job(
        project,
        document_code,
        tracking_document_sn=tracking_document_sn,
        initial_only=True,
        progress_codes=RUNNING_PROGRESS_CODES,
    )


def get_running_document(project, document_code, *, tracking_document_sn=None):
    return find_document_job(
        project,
        document_code,
        tracking_document_sn=tracking_document_sn,
        progress_codes=RUNNING_PROGRESS_CODES,
    )


def get_running_history_job(project, document_code):
    initial_document = get_running_initial_document(project, document_code)
    if initial_document is not None:
        return "initial", initial_document
    running_document = get_running_document(project, document_code)
    if running_document is not None:
        return "auto_apply", running_document
    return None, None


def start_initial_generation_job(project, actor, state):
    document_code = get_current_generation_code(state)
    if not document_code:
        return {"status": "error", "document": None, "message": "생성할 산출물 단계를 찾지 못했습니다."}

    running_document = get_running_initial_document(project, document_code)
    if running_document is not None:
        set_generation_draft_document(state, running_document)
        return {"status": "running", "document": running_document, "message": "문서를 생성 중입니다."}

    payload = build_generation_request_payload(project, state, document_code, update_mode="N")
    try:
        request_fastapi_generate(payload)
    except (HTTPError, URLError, ValueError) as exc:
        return {"status": "error", "document": None, "message": extract_fastapi_error_message(exc)}
    tracked_document = wait_for_document_job(project, document_code, initial_only=True)
    if tracked_document is None:
        return {"status": "error", "document": None, "message": "문서 생성 작업을 시작하지 못했습니다."}

    set_generation_draft_document(state, tracked_document)
    return {"status": "started", "document": tracked_document, "message": "문서 생성을 요청했습니다."}


def start_auto_apply_job(project, document_code, selected_files):
    running_document = get_running_document(project, document_code)
    if running_document is not None:
        return {"status": "running", "document": running_document, "message": "문서를 생성 중입니다."}

    payload = build_auto_apply_request_payload(project, document_code, selected_files)
    try:
        request_fastapi_generate(payload)
    except (HTTPError, URLError, ValueError) as exc:
        return {"status": "error", "document": None, "message": extract_fastapi_error_message(exc)}
    tracked_document = wait_for_document_job(project, document_code)
    if tracked_document is None:
        return {"status": "error", "document": None, "message": "회의 내용 자동 적용 작업을 시작하지 못했습니다."}
    return {"status": "started", "document": tracked_document, "message": "회의 내용 자동 적용을 요청했습니다."}


def build_generation_lines(project, document_code, inputs):
    label = get_document_label(document_code)
    if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE:
        rows = [
            f"{project.name} 프로젝트의 {label} 초안입니다.",
            "아래 업로드한 UI 참고 이미지를 기준으로 더미 자동 생성 결과를 구성했습니다.",
        ]
        for reference in inputs:
            rows.append(f"- {reference.get('name', '이미지 파일')}")
    elif document_code == ARCHITECTURE_DOCUMENT_CODE:
        rows = [
            f"{project.name} 프로젝트의 {label} 초안입니다.",
            "아래 등록된 서버 정보를 기준으로 더미 자동 생성 결과를 구성했습니다.",
        ]
        for project_net in inputs:
            description = project_net.purpose or "목적 미입력"
            rows.append(f"- {project_net.name} ({description})")
    else:
        rows = [
            f"{project.name} 프로젝트의 {label} 초안입니다.",
            "아래 선택한 문서를 기준으로 더미 자동 생성 결과를 구성했습니다.",
        ]
        for project_file in inputs:
            rows.append(f"- {project_file.name} ({project_file.file_type.name})")

    previous_document_code = get_previous_document_code(document_code)
    if previous_document_code:
        rows.append(
            f"이 문서는 직전 단계인 {get_document_label(previous_document_code)} 확정본을 이어받아 작성됩니다."
        )
    rows.append("내용을 검토한 뒤 OnlyOffice에서 수정하고 파일 확정을 진행해 주세요.")
    return rows


@transaction.atomic
def create_document_with_detail(
    *,
    project,
    document_code,
    actor,
    version,
    modification_content,
    content_bytes,
    locked_user=None,
    progress_status_id=PROGRESS_COMPLETED,
):
    try:
        document = Document.objects.create(
            project=project,
            possession_user=locked_user,
            document_type_id=document_code,
            progress_status_id=progress_status_id,
            version=version,
            modification_content=modification_content,
            created_by=actor,
            updated_by=actor,
        )
        detail = DocumentDetail.objects.create(
            document=document,
            path="",
            is_deleted="N",
            created_by=actor,
        )
        detail_path = build_document_detail_path(project, document.sn, detail.sn)
        save_bytes(
            build_document_detail_storage_key(project, document.sn, detail.sn),
            content_bytes or b"",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        detail.path = detail_path
        detail.save(update_fields=["path"])
    except Exception:
        if "detail_path" in locals():
            delete_object_at_uri(detail_path)
        raise
    return document, detail


def create_draft_document(project, document_code, actor, source_inputs):
    label = get_document_label(document_code)
    content_bytes = build_docx_bytes(label, build_generation_lines(project, document_code, source_inputs))
    return create_document_with_detail(
        project=project,
        document_code=document_code,
        actor=actor,
        version="0",
        modification_content="최초 생성",
        content_bytes=content_bytes,
        locked_user=None,
        progress_status_id=PROGRESS_PROCESSING,
    )


@transaction.atomic
def confirm_document(document, actor):
    latest_detail = get_latest_detail(document)
    confirmed_document, confirmed_detail = create_document_with_detail(
        project=document.project,
        document_code=document.document_type_id,
        actor=actor,
        version="1",
        modification_content="파일 확정",
        content_bytes=get_document_detail_bytes(latest_detail),
        locked_user=None,
        progress_status_id=PROGRESS_COMPLETED,
    )
    document.possession_user = None
    document.progress_status_id = PROGRESS_COMPLETED
    document.updated_by = actor
    document.save(update_fields=["possession_user", "progress_status", "updated_by"])
    return confirmed_document, confirmed_detail


@transaction.atomic
def acquire_document_lock(document, actor):
    if document.possession_user_id and document.possession_user_id != actor.sn:
        return False
    document.possession_user = actor
    document.updated_by = actor
    document.save(update_fields=["possession_user", "updated_by"])
    return True


@transaction.atomic
def release_document_lock(document, actor=None):
    document.possession_user = None
    if actor is not None:
        document.updated_by = actor
        document.save(update_fields=["possession_user", "updated_by"])
    else:
        document.save(update_fields=["possession_user"])


@transaction.atomic
def save_revision(document, actor, *, text_content=None, content_bytes=None, modification_content="수정 저장"):
    latest_detail = get_latest_detail(document)
    if content_bytes is None:
        if text_content is not None:
            content_bytes = build_docx_bytes(
                get_document_label(document.document_type_id),
                text_content.splitlines(),
            )
        elif latest_detail is not None:
            content_bytes = get_document_detail_bytes(latest_detail)

    try:
        detail = DocumentDetail.objects.create(
            document=document,
            path="",
            is_deleted="N",
            created_by=actor,
        )
        detail_path = build_document_detail_path(document.project, document.sn, detail.sn)
        save_bytes(
            build_document_detail_storage_key(document.project, document.sn, detail.sn),
            content_bytes or b"",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        detail.path = detail_path
        detail.save(update_fields=["path"])
    except Exception:
        if "detail_path" in locals():
            delete_object_at_uri(detail_path)
        raise
    document.updated_by = actor
    document.modification_content = modification_content
    document.save(update_fields=["updated_by", "modification_content"])
    return detail


def apply_meeting_notes(document, actor, selected_files):
    base_text = extract_text_from_docx(get_document_detail_bytes(get_latest_detail(document)))
    new_text = [
        base_text,
        "",
        "[회의 내용 자동 반영]",
        "다음 회의록을 반영한 더미 수정 결과입니다.",
    ]
    for project_file in selected_files:
        new_text.append(f"- {project_file.name}")
    return save_revision(
        document,
        actor,
        text_content="\n".join(new_text).strip(),
        modification_content="회의 내용 자동 반영",
    )


@transaction.atomic
def restore_revision(document, actor, source_detail):
    return save_revision(
        document,
        actor,
        content_bytes=get_document_detail_bytes(source_detail),
        modification_content="이전 버전 복원",
    )


def can_request_approval(document, actor, *, pending_approval=None, is_generation_draft=False):
    if actor is None or is_generation_draft or pending_approval is not None:
        return False
    return document.updated_by_id == actor.sn


@transaction.atomic
def create_approval_request(document, actor, request_content):
    latest_detail = get_latest_detail(document)
    approval = DocumentApproval.objects.create(
        detail=latest_detail,
        approval_status_id="APRV_REQ",
        request_content=request_content,
        rejection_reason=None,
        created_by=actor,
        updated_by=actor,
    )
    release_document_lock(document, actor)
    return approval


@transaction.atomic
def cancel_approval_request(approval):
    approval.delete()


def has_document_version(project, document_code, version):
    return Document.objects.filter(
        project=project,
        document_type_id=document_code,
        version=version,
    ).exists()


@transaction.atomic
def approve_request(approval, actor, new_version):
    source_detail = approval.detail
    source_document = source_detail.document
    document, detail = create_document_with_detail(
        project=source_document.project,
        document_code=source_document.document_type_id,
        actor=actor,
        version=new_version,
        modification_content=approval.request_content or "승인 반영",
        content_bytes=get_document_detail_bytes(source_detail),
        locked_user=None,
    )
    approval.approval_status_id = "APRV_COM"
    approval.updated_by = actor
    approval.save(update_fields=["approval_status", "updated_by"])
    return document, detail


@transaction.atomic
def reject_request(approval, actor, reason):
    approval.approval_status_id = "APRV_RJT"
    approval.rejection_reason = reason
    approval.updated_by = actor
    approval.save(update_fields=["approval_status", "rejection_reason", "updated_by"])


def get_generation_state(session, project):
    state = session.get(GENERATION_SESSION_KEY)
    if not state:
        return _build_empty_generation_state(project)
    if state.get("project_sn") != getattr(project, "sn", None):
        clear_generation_state(session)
        return _build_empty_generation_state(project)
    state.setdefault("selected_file_ids", [])
    state.setdefault("draft_documents", {})
    state.setdefault("confirmed_documents", {})
    state.setdefault("itf_reference_files", [])
    return state


def save_generation_state(session, state):
    session[GENERATION_SESSION_KEY] = state
    session.modified = True


def clear_generation_state(session, project=None):
    state = session.get(GENERATION_SESSION_KEY)
    if project is None or not state or state.get("project_sn") == getattr(project, "sn", None):
        if state:
            cleanup_generation_itf_references(state)
        session.pop(GENERATION_SESSION_KEY, None)
        session.modified = True


def update_generation_selected_files(state, file_ids):
    cleanup_generation_itf_references(state)
    state["selected_file_ids"] = [str(file_id) for file_id in file_ids if str(file_id).strip()]
    state["draft_documents"] = {}
    state["confirmed_documents"] = {}
    state["itf_reference_files"] = []
    return state


def get_generation_selected_files(project, state):
    return list(
        get_project_files(
            project,
            file_ids=state.get("selected_file_ids", []),
            allowed_types=ALLOWED_GENERATION_FILE_CODES,
        )
    )


def get_project_nets(project):
    if project is None:
        return []
    return list(ProjectNet.objects.filter(project=project).order_by("sn"))


def get_generation_source_inputs(project, state, document_code):
    if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE:
        return get_generation_itf_references(state)
    if document_code == ARCHITECTURE_DOCUMENT_CODE:
        return get_project_nets(project)
    return get_generation_selected_files(project, state)


def get_generation_prerequisite_error(project, state, document_code):
    if document_code == INTERFACE_REFERENCE_DOCUMENT_CODE and not get_generation_itf_references(state):
        return "사용자 인터페이스 참고 이미지를 하나 이상 업로드해 주세요."
    if document_code == ARCHITECTURE_DOCUMENT_CODE and not get_project_nets(project):
        return "서버 정보를 하나 이상 추가해 주세요."
    if document_code not in {INTERFACE_REFERENCE_DOCUMENT_CODE, ARCHITECTURE_DOCUMENT_CODE} and not get_generation_selected_files(project, state):
        return "생성에 사용할 문서를 먼저 선택해 주세요."
    return None


def get_current_generation_code(state):
    confirmed = state.get("confirmed_documents", {})
    for code in get_document_code_sequence():
        if str(code) not in confirmed:
            return code
    return None


def is_generation_complete(state):
    return get_current_generation_code(state) is None


def get_generation_progress_rows(state):
    rows = []
    current_code = get_current_generation_code(state)
    for code in get_document_code_sequence():
        if code in state.get("confirmed_documents", {}):
            status = "confirmed"
            status_label = "확정 완료"
        elif code in state.get("draft_documents", {}):
            status = "review"
            status_label = "검토 중"
        elif current_code == code:
            status = "pending"
            status_label = "생성 대기"
        else:
            status = "locked"
            status_label = "이전 단계 대기"
        rows.append(
            {
                "code": code,
                "label": get_document_label(code),
                "status": status,
                "status_label": status_label,
            }
        )
    return rows


def ensure_generation_draft(project, actor, state):
    document_code = get_current_generation_code(state)
    if not document_code:
        return None, False

    existing_sn = state.get("draft_documents", {}).get(document_code)
    if existing_sn:
        existing_document = (
            Document.objects.filter(
                sn=existing_sn,
                project=project,
                document_type_id=document_code,
                version="0",
            )
            .select_related("project", "document_type", "created_by", "updated_by", "possession_user")
            .first()
        )
        if existing_document is not None:
            return existing_document, False

    source_inputs = get_generation_source_inputs(project, state, document_code)
    if not source_inputs:
        return None, False

    draft_document, _ = create_draft_document(project, document_code, actor, source_inputs)
    state.setdefault("draft_documents", {})[document_code] = draft_document.sn
    return draft_document, True


def mark_generation_confirmed(state, draft_document, confirmed_document):
    code = draft_document.document_type_id
    state.setdefault("confirmed_documents", {})[code] = confirmed_document.sn
    state.setdefault("draft_documents", {}).pop(code, None)
    return state


def build_generation_redirect_url(*, document_code=None, play=False, auto_start=False, resume=False):
    query_items = []
    if document_code:
        query_items.append(("docs_cd", document_code))
    if resume:
        query_items.append(("resume", "1"))
    if play:
        query_items.append(("play", "1"))
    if auto_start:
        query_items.append(("auto_start", "1"))
    base_url = reverse("doc_generate")
    return f"{base_url}?{urlencode(query_items)}" if query_items else base_url


def build_history_preview_url(document, preview_detail_sn=None, *, mode=None):
    query_items = [("docs_cd", document.document_type_id)]
    if preview_detail_sn is not None:
        query_items.append(("preview_detail", str(preview_detail_sn)))
        query_items.append(("modal", "history"))
    if mode == "edit":
        query_items.append(("mode", "edit"))
    return f"{reverse('doc_detail', args=[document.sn])}?{urlencode(query_items)}"


def build_history_preview_api_url(document, detail_sn):
    return reverse("doc_history_preview", args=[document.sn, detail_sn])


def build_document_detail_url(document, *, mode=None):
    query_items = []
    if mode:
        query_items.append(("mode", mode))
    base_url = reverse("doc_detail", args=[document.sn])
    return f"{base_url}?{urlencode(query_items)}" if query_items else base_url


def build_generation_steps(document_code):
    label = get_document_label(document_code)
    return [
        f"{label} 자동 생성을 요청했습니다.",
        "선택한 파일을 분석하고 더미 응답을 준비하고 있습니다.",
        "초안 구조를 정리하고 있습니다.",
        f"{label} 초안이 준비되었습니다. 확인 버튼으로 결과를 검토해 주세요.",
    ]


@transaction.atomic
def create_project_net(
    *,
    project,
    actor,
    name,
    purpose="",
    middleware_stack="",
    firewall_settings="",
    auth_method="",
    expected_concurrent_users=None,
    cloud_yn=None,
    hardware_spec="",
    remarks="",
):
    return ProjectNet.objects.create(
        project=project,
        name=name,
        purpose=purpose or None,
        middleware_stack=middleware_stack or None,
        firewall_settings=firewall_settings or None,
        auth_method=auth_method or None,
        expected_concurrent_users=expected_concurrent_users,
        cloud_yn=cloud_yn,
        hardware_spec=hardware_spec or None,
        remarks=remarks or None,
        created_by=actor,
        updated_by=actor,
    )


def get_document_view_state(document, actor, preferred_mode="view"):
    pending_approval = get_latest_pending_approval(document)
    if pending_approval and pending_approval.created_by_id == actor.sn:
        return "waiting", pending_approval

    if document.possession_user_id and document.possession_user_id != actor.sn:
        return "readonly", pending_approval
    if document.possession_user_id and document.possession_user_id == actor.sn and preferred_mode == "edit":
        return "edit", pending_approval
    return "view", pending_approval


def wait_for_new_revision(document, *, baseline_detail_sn=None, timeout_seconds=5, interval_seconds=0.25):
    deadline = time.monotonic() + timeout_seconds
    while time.monotonic() < deadline:
        latest_detail = get_latest_detail(document)
        if latest_detail is not None and latest_detail.sn != baseline_detail_sn:
            return latest_detail
        time.sleep(interval_seconds)
    return get_latest_detail(document)


def build_editor_config(request, document, actor, mode):
    latest_detail = get_latest_detail(document)
    public_base_url = getattr(settings, "DJANGO_PUBLIC_BASE_URL", "").rstrip("/")
    if not public_base_url:
        public_base_url = request.build_absolute_uri("/").rstrip("/")

    document_query = ""
    if settings.ONLYOFFICE_JWT_SECRET:
        document_query = urlencode(
            {
                "token": encode_jwt(
                    {"document_sn": document.sn, "project_sn": document.project_id},
                    settings.ONLYOFFICE_JWT_SECRET,
                )
            }
        )

    document_url = f"{public_base_url}{reverse('doc_content', args=[document.sn])}"
    if document_query:
        document_url = f"{document_url}?{document_query}"

    payload = {
        "documentType": "word",
        "width": "100%",
        "document": {
            "title": get_document_title(document),
            "url": document_url,
            "fileType": "docx",
            "key": build_document_key(document, latest_detail=latest_detail),
            "permissions": {
                "edit": mode == "edit",
                "download": True,
                "print": True,
                "comment": False,
                "review": False,
            },
        },
        "editorConfig": {
            "callbackUrl": f"{public_base_url}{reverse('doc_callback', args=[document.sn])}",
            "mode": mode,
            "user": {"id": str(actor.sn), "name": actor.name},
        },
        "type": "desktop" if mode == "edit" else "embedded",
    }

    if settings.ONLYOFFICE_JWT_SECRET:
        payload["token"] = encode_jwt(payload, settings.ONLYOFFICE_JWT_SECRET)
    return payload


def parse_callback_payload(request):
    body = request.body.decode("utf-8") if request.body else "{}"
    payload = json.loads(body)

    auth_header = request.headers.get("Authorization", "")
    token = None
    if auth_header.startswith("Bearer "):
        token = auth_header.split(" ", 1)[1].strip()
    elif payload.get("token"):
        token = payload["token"]

    if settings.ONLYOFFICE_JWT_SECRET and token:
        decode_jwt(token, settings.ONLYOFFICE_JWT_SECRET)
    return payload


def validate_document_content_token(document, token):
    if not settings.ONLYOFFICE_JWT_SECRET or not token:
        return False
    try:
        payload = decode_jwt(token, settings.ONLYOFFICE_JWT_SECRET)
    except Exception:
        return False
    return (
        str(payload.get("document_sn")) == str(document.sn)
        and str(payload.get("project_sn")) == str(document.project_id)
    )


def build_consistency_review(approval):
    requester = approval.created_by.name if approval.created_by else "작성자"
    return {
        "title": "정합성 자동 검토 결과",
        "summary": f"{requester}의 수정 요청을 검토했습니다.",
        "items": [
            "수정본과 직전 버전 간 문서 구조 차이는 정상 범위입니다.",
            "회의록 또는 RFP에서 추출한 변경 사유와 충돌하는 항목은 발견되지 않았습니다.",
            "최종 승인 전 오탈자와 버전명만 다시 확인하면 됩니다.",
        ],
    }


def get_project_files(project, *, file_ids=None, allowed_types=None):
    queryset = ProjectFile.objects.filter(project=project).select_related("file_type", "created_by")
    if allowed_types:
        queryset = queryset.filter(file_type_id__in=allowed_types)
    if file_ids:
        queryset = queryset.filter(sn__in=file_ids)
    return queryset.order_by("-created_at", "-sn")


def build_document_rows(queryset):
    rows = []
    for document in queryset:
        rows.append(
            {
                "sn": document.sn,
                "type_name": getattr(document.document_type, "name", "-") or "-",
                "creator_name": getattr(document.created_by, "name", "-") or "-",
                "version": document.version or "-",
                "modification_content": document.modification_content or "-",
                "created_at": document.created_at,
                "detail_url": reverse("doc_detail", args=[document.sn]),
                "locked_by_name": getattr(document.possession_user, "name", ""),
            }
        )
    return rows


def build_approval_rows(queryset):
    rows = []
    for approval in queryset:
        rows.append(
            {
                "sn": approval.approval_sn,
                "document_sn": approval.detail.document.sn,
                "document_label": approval.detail.document.document_type.name,
                "version": approval.detail.document.version,
                "requester_name": getattr(approval.created_by, "name", "-") or "-",
                "status_name": approval.approval_status.name,
                "request_content": approval.request_content,
                "created_at": approval.created_at,
                "detail_url": reverse("doc_approval_detail", args=[approval.approval_sn]),
            }
        )
    return rows


def build_approval_queryset(project, actor):
    queryset = (
        DocumentApproval.objects.filter(detail__document__project=project)
        .select_related(
            "detail__document__document_type",
            "approval_status",
            "created_by",
            "detail__document__project",
        )
        .order_by("-created_at", "-approval_sn")
    )
    if not is_project_manager(project, actor):
        queryset = queryset.filter(created_by=actor)
    return queryset


def apply_approval_filters(params, queryset, *, include_requester=True):
    document_code = params.get("docs_cd", "all")
    approval_status = params.get("status", "all")
    requester_query = params.get("requester", "").strip()

    if document_code != "all":
        queryset = queryset.filter(detail__document__document_type_id=document_code)
    if approval_status != "all":
        queryset = queryset.filter(approval_status_id=approval_status)
    if include_requester and requester_query:
        queryset = queryset.filter(created_by__name__icontains=requester_query)

    return queryset, document_code, approval_status, requester_query
